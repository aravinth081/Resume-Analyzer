"""
Resume Parser — Extracts raw text from PDF and DOCX files.

Uses pdfplumber for PDFs (better table/layout extraction than PyPDF2)
and python-docx for DOCX files.
"""
import os
import re
from typing import Optional
import pdfplumber
from docx import Document


def parse_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using pdfplumber.
    Handles multi-page documents and preserves paragraph structure.
    """
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    Processes paragraphs and tables to capture all content.
    """
    doc = Document(file_path)
    text_parts = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract table content (resumes sometimes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def parse_resume_file(file_path: str) -> str:
    """
    Parse a resume file (PDF or DOCX) and return the extracted raw text.
    Raises ValueError for unsupported file types.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .pdf and .docx are supported.")


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted resume text.
    - Remove excessive whitespace
    - Normalize line breaks
    - Remove special characters that break NLP processing
    """
    # Replace multiple newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    # Remove non-printable characters (keep newlines and tabs)
    text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(lines).strip()
